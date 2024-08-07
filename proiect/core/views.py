from django.contrib.auth import authenticate, login
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.forms import AuthenticationForm
from django.contrib.auth.decorators import login_required
from proiect.forms import ReviewerSignupForm, DocumentUploadForm , ConferenceForm
from proiect.models import UploadedDocument, ReviewerRequest, Conference, CustomUser, ReviewFeedback
from django.contrib.auth import get_user_model
from django.core.mail import send_mail
from django.http import JsonResponse
from django.contrib.auth import get_user_model
from docx import Document
from django.http import HttpResponse
from django.db.models import Q
from django.shortcuts import redirect
from os.path import basename
from gensim import corpora, similarities
from transformers import BertTokenizer, BertForSequenceClassification
from wordcloud import WordCloud
from datetime import date
from django.contrib import messages
import string
from gensim.parsing.preprocessing import remove_stopwords
from nltk.stem import WordNetLemmatizer
import nltk
import numpy as np

import matplotlib.pyplot as plt
import fitz
import enchant
import io
import base64
import torch
import spacy
import requests

nltk.download('wordnet')

nlp = spacy.load("en_core_web_sm")


def index(request):
    country_name, country_code = get_country_from_ip(request)
    
    if country_name:
        all_conferences = Conference.objects.all()
        conferences_in_country = 0
        for conference in all_conferences:
            conference_country = get_country_from_location(conference.location)
            if conference_country and conference_country == country_name:
                conferences_in_country += 1
    else:
        country_name = "Your Location"
        country_code = ""
        conferences_in_country = 0
    
    total_conferences = Conference.objects.all().count()

    context = {
        "title": "Home",
        'country': country_name,
        'country_flag': country_code,
        'conferences_in_country': conferences_in_country,
        'total_conferences': total_conferences,
    }
    return render(request, "main_templates/index.html", context)

def get_country_from_ip(request):
    try:
        response = requests.get('https://ipinfo.io/json')
        data = response.json()
        country_code = data['country']
        country_name = COUNTRY_CODES.get(country_code, None)
        return country_name, country_code
    except Exception as e:
        return None, None
    
def get_country_from_location(location):
    parts = location.split(',')
    if parts:
        return parts[0].strip()
    return None

COUNTRY_CODES = {
    'RO': 'Romania',
    'US': 'United States',
    'PL': 'Poland',
}

#user
@login_required
def user_dash(request):
    documents = UploadedDocument.objects.filter(user=request.user)
    conferences = Conference.objects.all()

    for document in documents:
        rejected_document = UploadedDocument.objects.filter(document=document.document.name, status='REJECTED').first()
        if rejected_document and rejected_document.conference:
            document.rejected_conference = rejected_document.conference.name
        else:
            document.rejected_conference = None


    if request.method == 'POST':
        form = DocumentUploadForm(request.POST, request.FILES)
        if form.is_valid():
            new_document = form.save(commit=False)
            new_document.user = request.user
            new_document.save()
            return redirect('user_dash')
    else:
        form = DocumentUploadForm()

    return render(request, 'dashboard_templates/user_dashboard.html', {
        'form': form,
        'documents': documents,
        'conferences': conferences,
    })

def get_feedback(request, document_id):
    feedback = ReviewFeedback.objects.filter(document_id=document_id)

    feedback_html = ""
    for entry in feedback:
        feedback_html += f"<p><b>What's Wrong:</b> <br> {entry.whats_wrong}</p>"
        feedback_html += f"<p><b>What Can Be Improved:</b> <br> {entry.what_can_be_improved}</p>"
        feedback_html += f"<p><b>Score:</b> {entry.score}</p>"
        feedback_html += f"<p><b>Decision:</b> {entry.get_decision_display()}</p>"
        feedback_html += "<hr>"

    return render(request, 'results_templates/feedback.html', {'feedback_html': feedback_html})

@login_required
def remove_document(request, document_id):
    document = get_object_or_404(UploadedDocument, id=document_id, user=request.user)
    document.delete()
    return redirect(request.META.get('HTTP_REFERER', ''))

@login_required
def send_to_review(request, document_id):
    document = get_object_or_404(UploadedDocument, id=document_id, user=request.user)

    if request.method == 'POST':
        conference_name = request.POST.get('conference_name')
        keywords = request.POST.get('keywords')
        topic = request.POST.get('topic')

        document.keywords = keywords
        document.topic = topic
        document.status = 'SUBMITTED'
        document.save()

        conference = get_object_or_404(Conference, name=conference_name)
        trackers = conference.trackers.filter(is_tracker=True)

        document.conference = conference
        document.save()
        
        for tracker in trackers:
            UploadedDocument.objects.create(
                user=tracker,
                document=document.document,
                keywords=keywords,
                topic=topic,
                status='SUBMITTED',
                reviewer=tracker,
                conference=conference
            )

        messages.success(request, 'Document has been submitted for review.')
        return redirect('user_dash')

    return render(request, 'dashboard_templates/user_dashboard.html', {
        'document': document,
        'form': DocumentUploadForm(),
    })


def preview_document(request, document_id):
    try:
        document = UploadedDocument.objects.get(id=document_id)
        file_path = document.document.path
        if file_path.endswith('.pdf'):
            with open(file_path, 'rb') as f:
                response = HttpResponse(f.read(), content_type='application/pdf')
            return response
        elif file_path.endswith('.docx'):
            doc = Document(file_path)
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            return HttpResponse(text, content_type='text/plain; charset=utf-8')
        else:
            return HttpResponse("Unsupported file format")
    except UploadedDocument.DoesNotExist:
        return HttpResponse("Document not found", status=404)

def spell_check_view(request):
    if request.method == 'POST':
        document_id = request.POST.get('document_id')
        try:
            document = UploadedDocument.objects.get(id=document_id)
            file_path = document.document.path

            if file_path.endswith('.pdf'):
                text = extract_text_from_pdf(file_path)
            elif file_path.endswith('.docx'):
                text = extract_text_from_docx(file_path)
            else:
                return render(request, 'error.html', {'message': 'Unsupported file format'})

            mistakes = spell_check(text)

            return render(request, 'results_templates/spell_check_results.html', {'text': text, 'mistakes': mistakes})
        except UploadedDocument.DoesNotExist:
            return render(request, 'error.html', {'message': 'Document not found'})

def extract_text_from_pdf(pdf_path):
    text = ''
    with fitz.open(pdf_path) as doc:
        for page in doc:
            text += page.get_text()
    return text

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    return text

def spell_check(text):
    d = enchant.Dict("en_US")
    mistakes = []
    for word in text.split():
        if not d.check(word):
            mistakes.append(word)
    return mistakes

def ner_view(request):
    if request.method == 'POST':
        document_id = request.POST.get('document_id')
        try:
            document = UploadedDocument.objects.get(id=document_id)
            file_path = document.document.path

            if file_path.endswith('.pdf'):
                text = extract_text_from_pdf(file_path)
            elif file_path.endswith('.docx'):
                text = extract_text_from_docx(file_path)
            else:
                return render(request, 'error.html', {'message': 'Unsupported file format'})

            entities = ner_extraction(text)

            present_labels = set(entity[1] for entity in entities)

            word_cloud_img = generate_word_cloud(entities)

            category_labels = [
                ('ORG', 'Organizations'),
                ('PERSON', 'Persons'),
                ('NORP', 'Nationalities, Religious, Political Groups'),
                ('GPE', 'Countries, Cities, States'),
                ('LOC', 'Locations'),
                ('PRODUCT', 'Products'),
                ('EVENT', 'Events'),
                ('WORK_OF_ART', 'Artworks'),
                ('LAW', 'Laws'),
                ('LANGUAGE', 'Languages')
            ]

            return render(request, 'results_templates/ner_results.html', {
                'word_cloud_img': word_cloud_img,
                'entities': entities,
                'present_labels': present_labels,
                'category_labels': category_labels
            })
        except UploadedDocument.DoesNotExist:
            return render(request, 'error.html', {'message': 'Document not found'})

def ner_extraction(text):
    doc = nlp(text)
    
    entities = [(ent.text, ent.label_) for ent in doc.ents]
    return entities

def generate_word_cloud(entities):
    text = ' '.join([entity[0] for entity in entities])
    
    wordcloud = WordCloud(width=800, height=400, background_color='white').generate(text)

    plt.figure(figsize=(10, 5))
    plt.imshow(wordcloud, interpolation='bilinear')
    plt.axis('off')
    
    buffer = io.BytesIO()
    plt.savefig(buffer, format='png')
    buffer.seek(0)
    
    img_str = base64.b64encode(buffer.read()).decode()
    buffer.close()
    
    return img_str



def analyze_sentiment_bert(text):
    tokenizer = BertTokenizer.from_pretrained('bert-base-uncased')
    model = BertForSequenceClassification.from_pretrained('bert-base-uncased')

    inputs = tokenizer(text, return_tensors='pt', padding=True, truncation=True)
    
    outputs = model(**inputs)

    predicted_class = torch.argmax(outputs.logits, dim=1).item()
    sentiment = ['negative', 'neutral', 'positive'][predicted_class]

    return sentiment

def analyze_sentiment_view(request):
    if request.method == 'POST':
        document_id = request.POST.get('document_id')
        try:
            document = UploadedDocument.objects.get(id=document_id)
            file_path = document.document.path

            if file_path.endswith('.pdf'):
                text = extract_text_from_pdf(file_path)
            elif file_path.endswith('.docx'):
                text = extract_text_from_docx(file_path)
            else:
                return JsonResponse({'error': 'Unsupported file format'}, status=400)

            sentiment = analyze_sentiment_bert(text)

            return JsonResponse({'sentiment': sentiment})
        except UploadedDocument.DoesNotExist:
            return JsonResponse({'error': 'Document not found'}, status=404)

def compare_documents(selected_document_id):
    try:
        selected_document = UploadedDocument.objects.get(id=selected_document_id)
        selected_document_text = get_document_text(selected_document)

        other_documents = UploadedDocument.objects.filter(~Q(id=selected_document_id), reviewer__isnull=True)
        other_document_texts = [get_document_text(doc) for doc in other_documents]

        processed_texts = [preprocess_text(selected_document_text)] + [preprocess_text(text) for text in other_document_texts]

        dictionary = corpora.Dictionary(processed_texts)
        corpus = [dictionary.doc2bow(text) for text in processed_texts]

        index = similarities.MatrixSimilarity(corpus)

        sims = index[corpus[0]]
        
        sims = np.delete(sims, 0)

        sorted_sims = sorted(enumerate(sims), key=lambda item: -item[1])

        top_similar_documents = [(other_documents[i], sim) for i, sim in sorted_sims[:2]]

        return top_similar_documents

    except UploadedDocument.DoesNotExist:
        return None

def get_document_text(document):
    if document.document.path.endswith('.pdf'):
        return extract_text_from_pdf(document.document.path)
    elif document.document.path.endswith('.docx'):
        return extract_text_from_docx(document.document.path)
    else:
        return ""

def compare_documents_view(request):
    if request.method == 'POST':
        document_id = request.POST.get('document_id')
        similar_documents = compare_documents(document_id)
        if similar_documents is not None:
            return render(request, 'results_templates/similarity_results.html', {'similar_documents': similar_documents})
        else:
            return render(request, 'error.html', {'message': 'No similar documents found.'})
    else:
        return render(request, 'error.html', {'message': 'Invalid request method'})

def preprocess_text(text):
    text = text.lower()
    text = text.translate(str.maketrans('', '', string.punctuation))
    text = remove_stopwords(text)
    lemmatizer = WordNetLemmatizer()
    text = ' '.join([lemmatizer.lemmatize(word) for word in text.split()])
    return text.split()


#reviewer
@login_required
def reviewer_dash(request):
    reviewer = request.user
    submitted_documents = UploadedDocument.objects.filter(user=reviewer, status='SUBMITTED')
    reviewing_documents = UploadedDocument.objects.filter(user=reviewer, status='UNDER_REVIEW')
    
    search_query = request.GET.get('search_query')
    search_keyword = request.GET.get('search_keyword')
    search_topic = request.GET.get('search_topic')

    if search_query:
        query_filter = Q()

        if search_keyword and search_topic:
            query_filter |= Q(keywords__icontains=search_query) | Q(topic__icontains=search_query)
        elif search_keyword:
            query_filter |= Q(keywords__icontains=search_query)
        elif search_topic:
            query_filter |= Q(topic__icontains=search_query)
        else:
            query_filter |= Q(keywords__icontains=search_query) | Q(topic__icontains=search_query)

        submitted_documents = submitted_documents.filter(query_filter)

        if not submitted_documents.exists():
            no_document_message = "No document matched the query."
        else:
            no_document_message = None
    else:
        search_query = None
        no_document_message = None

    joined_conferences = reviewer.joined_conferences.all()

    return render(request, "dashboard_templates/reviewer_dashboard.html", {
        'submitted_documents': submitted_documents,
        'reviewing_documents': reviewing_documents,
        'search_query': search_query,
        'no_document_message': no_document_message,
        'reviewer': reviewer,
        'joined_conferences': joined_conferences
    })


@login_required
def choose_document(request, document_id):
    document = get_object_or_404(UploadedDocument, id=document_id, user=request.user, status='SUBMITTED')
    document.status = 'UNDER_REVIEW'
    document.save()
    document_name = basename(document.document.name)
    
    documents_with_same_name = UploadedDocument.objects.filter(document__icontains=document_name)
    
    for document in documents_with_same_name:
        if not document.reviewer:
            document.status = 'UNDER_REVIEW'
            document.save()
    
    return redirect('reviewer_dash')

@login_required
def return_document(request, document_id):
    document = get_object_or_404(UploadedDocument, id=document_id, user=request.user, status='UNDER_REVIEW')
    uploader_document = UploadedDocument.objects.filter(document=document.document.name, reviewer=None).first()

    if request.method == 'POST':
        whats_wrong = request.POST.get('whats_wrong')
        what_can_be_improved = request.POST.get('what_can_be_improved')
        score = request.POST.get('score')
        decision = request.POST.get('decision')

        feedback = ReviewFeedback.objects.create(
            document=uploader_document,
            reviewer=request.user,
            whats_wrong=whats_wrong,
            what_can_be_improved=what_can_be_improved,
            score=score,
            decision=decision
        )

        document.status = 'REVIEWED'
        document.feedback = feedback
        document.save()

        if UploadedDocument.objects.filter(
             conference=document.conference,
             document=document.document.name,
             status='UNDER_REVIEW',
             reviewer__isnull=False
        ).exclude(id=document.id).exists():
         uploader_document.status = 'UNDER_REVIEW'
        else:
         uploader_document.status = 'REVIEWED'

        uploader_document.feedback = feedback
        uploader_document.save()

        return redirect('reviewer_dash')

    return redirect('reviewer_dash')


#tracker    
@login_required
def tracker_dash(request):
    documents = UploadedDocument.objects.filter(reviewer=request.user)

    for document in documents:
        uploader_document = UploadedDocument.objects.filter(document=document.document.name, reviewer=None).first()
        if uploader_document:
            document.uploader = uploader_document.user.username
            document.uploader_workplace = uploader_document.user.current_workplace
            document.uploader_references = uploader_document.user.references
        else:
            document.uploader = "Unknown"
            document.uploader_workplace = "Unknown"
            document.uploader_references = "Unknown"

    tracked_conferences = Conference.objects.filter(trackers=request.user)

    reviewers = CustomUser.objects.filter(is_reviewer=True, joined_conferences__in=tracked_conferences)

    ongoing_reviews = UploadedDocument.objects.filter(
        status='UNDER_REVIEW',
        reviewer__isnull=False,
        conference__in=tracked_conferences
    ).distinct()

    grouped_ongoing_reviews = {}
    for review in ongoing_reviews:
        if review.document.name not in grouped_ongoing_reviews:
            grouped_ongoing_reviews[review.document.name] = []
        grouped_ongoing_reviews[review.document.name].append(review)

    return render(request, 'dashboard_templates/tracker_dashboard.html', {
        'documents': documents,
        'tracked_conferences': tracked_conferences,
        'reviewers': reviewers,
        'grouped_ongoing_reviews': grouped_ongoing_reviews,
    })



def reject_document(request, document_id):
    document = UploadedDocument.objects.get(pk=document_id)
    
    document.status = 'REJECTED'
    document.save()
    
    related_document = UploadedDocument.objects.filter(document=document.document.name, reviewer=None).first()
    if related_document:
        related_document.status = 'UPLOADED'
        related_document.save()
    
    return redirect('tracker_dash')

@login_required
def match_reviewer(request, document_id):
    if request.method == 'POST':
        document = get_object_or_404(UploadedDocument, id=document_id)
        reviewer_id = request.POST.get('reviewer')

        reviewer = get_object_or_404(CustomUser, id=reviewer_id)

        UploadedDocument.objects.create(
            user=reviewer,
            document=document.document,
            keywords=document.keywords,
            topic=document.topic,
            status='SUBMITTED',
            reviewer=reviewer,
            conference=document.conference
        )
        
        return redirect('tracker_dash')

#organizer
@login_required
def organizer_dash(request):
    conferences = Conference.objects.filter(organizer=request.user)
    search_query = request.GET.get('search_query', '')
    users = CustomUser.objects.filter(
        Q(username__icontains=search_query) | Q(email__icontains=search_query)
    ) if search_query else None
    
    return render(request, 'dashboard_templates/organizer_dashboard.html', {
        'conferences': conferences,
        'users': users,
    })


@login_required
def assign_user_to_conference(request):
    if request.method == 'POST':
        for key, value in request.POST.items():
            if key.startswith('role_'):
                user_id = key.split('_')[1]
                role = value
                conference_id = request.POST.get(f'conference_{user_id}')
                if conference_id:
                    user = get_object_or_404(CustomUser, id=user_id)
                    conference = get_object_or_404(Conference, id=conference_id)
                    
                    if role == 'reviewer':
                        if conference in user.joined_conferences.all():
                            messages.error(request, f'{user.username} is already a reviewer for this conference.')
                        elif user in conference.trackers.all():
                            messages.error(request, f'A user can\'t be both reviewer and tracker for the same conference.')
                        else:
                            user.joined_conferences.add(conference)
                            user.save()
                            messages.success(request, f'{user.username} has been assigned as a reviewer for {conference.name}.')
                    
                    elif role == 'tracker':
                        if user in conference.trackers.all():
                            messages.error(request, f'{user.username} is already a tracker for this conference.')
                        elif conference in user.joined_conferences.all():
                            messages.error(request, f'A user can\'t be both tracker and reviewer for the same conference.')
                        else:
                            conference.trackers.add(user)
                            if not user.is_tracker:
                                user.is_tracker = True
                                user.save()
                            conference.save()
                            messages.success(request, f'{user.username} has been assigned as a tracker for {conference.name}.')
        
        return redirect('organizer_dash')
    return HttpResponse(status=405)

@login_required
def create_conference(request):
    if request.method == 'POST':
        form = ConferenceForm(request.POST, request.FILES)
        if form.is_valid():
            conference = form.save(commit=False)
            conference.organizer = request.user
            conference.save()
            return redirect('organizer_dash')
    else:
        form = ConferenceForm()
    return render(request, 'dashboard_templates/organizer_dashboard.html', {'form': form})

@login_required
def delete_conference(request, conference_id):
    conference = get_object_or_404(Conference, id=conference_id, organizer=request.user)
    if request.method == 'POST':
        conference.delete()
        return redirect('organizer_dash')
    return HttpResponse(status=405)

#admin
@login_required
def admin_dash(request):
    signup_requests = ReviewerRequest.objects.all()
    ongoing_reviews = UploadedDocument.objects.filter(status='UNDER_REVIEW', reviewer=None).select_related('conference', 'user')
    ongoing_reviewers_dict = {}

    search_query = request.GET.get('search_query')
    if search_query:
        ongoing_reviews = ongoing_reviews.filter(document__icontains=search_query)
    
    for document in ongoing_reviews:
        ongoing_reviewers = UploadedDocument.objects.filter(
            document__icontains=document.document.name,
            reviewer__isnull=False,
            status='UNDER_REVIEW'
        ).values_list('reviewer__username', 'reviewer__email').distinct()
        ongoing_reviewers_dict[document] = ongoing_reviewers
    
    user_search_query = request.GET.get('user_search_query')
    users = []
    if user_search_query:
        users = get_user_model().objects.filter(username__icontains=user_search_query)
    
    return render(request, 'dashboard_templates/admin_dashboard.html', {
        'signup_requests': signup_requests,
        'ongoing_reviews': ongoing_reviews,
        'ongoing_reviewers_dict': ongoing_reviewers_dict,
        'search_query': search_query,
        'user_search_query': user_search_query,
        'users': users,
    })


@login_required
def change_user_roles(request):
    if request.method == 'POST':
        for user in get_user_model().objects.all():
            user_id = request.POST.get('user_id')
            is_reviewer = request.POST.get(f'is_reviewer_{user_id}') == 'on'
            is_tracker = request.POST.get(f'is_tracker_{user_id}') == 'on'
            is_organizer = request.POST.get(f'is_organizer_{user_id}') == 'on'

            if (user.id == int(user_id)) and ((user.is_reviewer != is_reviewer) or (user.is_tracker != is_tracker) or (user.is_organizer != is_organizer)):
                user.is_reviewer = is_reviewer
                user.is_tracker = is_tracker
                user.is_organizer = is_organizer
                user.save()
    return redirect('admin_dash')


def approve_signup_request(request, request_id):
    signup_request = get_object_or_404(ReviewerRequest, id=request_id)
    
    user = get_user_model().objects.create_user(
        username=signup_request.username,
        email=signup_request.email,
        password=signup_request.password,
        current_workplace=signup_request.current_workplace,
        references=signup_request.references,
        is_active=True,
        is_reviewer=True
    )

    signup_request.delete()

    return redirect('admin_dash')

def deny_signup_request(request, request_id):
    signup_request = get_object_or_404(ReviewerRequest, id=request_id)
    signup_request.delete()
    return redirect('admin_dash')



def login_view(request):
    if request.method == 'POST':
        form = AuthenticationForm(data=request.POST)
        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            user = authenticate(request, username=username, password=password)
            if user is not None:
                login(request, user)
                return redirect('/')
    else:
        form = AuthenticationForm()

    return render(request, 'auth_templates/login.html', {'form': form})

def signup_user(request):
    if request.method == 'POST':
        form = ReviewerSignupForm(request.POST)
        if form.is_valid():
            form.save()
    else:
        form = ReviewerSignupForm(initial={'is_reviewer': False})



    return render(request, 'auth_templates/signup_user.html', {'form': form})

def signup_reviewer(request):
    if request.method == 'POST':
        form = ReviewerSignupForm(request.POST)
        if form.is_valid():
            email = form.cleaned_data['email']
            username = form.cleaned_data['username']
            password = form.cleaned_data['password1']


            ReviewerRequest.objects.create(
                email=email,
                username=username,
                current_workplace=form.cleaned_data['current_workplace'],
                references=form.cleaned_data['references'],
                password=password
            )

            return redirect('signup_reviewer_confirmation')
    else:
        form = ReviewerSignupForm()

    return render(request, 'auth_templates/signup_reviewer.html', {'form': form})

def signup_reviewer_confirmation(request):

    return render(request, "auth_templates/signup_reviewer_confirmation.html")

def your_view(request):
    user = request.user

    return render(request, 'main_templates/base.html', {'user': user})


#help

def help_page(request):
    return render(request, 'main_templates/help.html')

def contact_form(request):
    if request.method == 'POST':
        name = request.POST.get('title')
        email = request.POST.get('email')
        message = request.POST.get('message')

        send_mail(
            'Contact Form Submission',
            f'Title: {name}\nEmail: {email}\nMessage: {message}',
            email,
            ['mihaixz4@gmail.com'],
            fail_silently=False,
        )

        return JsonResponse({'status': 'success'})
    
# conferences 

def conferences(request):
    today = date.today()
    ongoing_conferences = Conference.objects.filter(start_date__lte=today, end_date__gte=today)
    planned_conferences = Conference.objects.filter(start_date__gt=today)
    return render(request, 'main_templates/conferences.html', {'ongoing_conferences': ongoing_conferences, 'planned_conferences': planned_conferences})
